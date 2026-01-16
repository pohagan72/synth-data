import os
import uuid
import json
import yaml
import glob
import random
import re
import time
import base64
import shutil
import zipfile
import tarfile
import threading
from io import BytesIO
from email.message import EmailMessage
from datetime import datetime, timedelta, timezone
from openai import AzureOpenAI
from dotenv import load_dotenv
from docx import Document
from openpyxl import Workbook
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- Configuration and Setup ---
load_dotenv()

def select_config_file():
    """Scans for .yaml files and prompts the user to select one."""
    try:
        yaml_files = [f for f in glob.glob('*.yaml') if 'requirements' not in f]
        if not yaml_files:
            print("Error: No .yaml configuration files found.")
            return None
        print("Please select a configuration file to use:")
        for i, filename in enumerate(yaml_files): print(f"  {i + 1}: {filename}")
        while True:
            try:
                choice = int(input(f"Enter a number (1-{len(yaml_files)}): ")) - 1
                if 0 <= choice < len(yaml_files):
                    selected_file = yaml_files[choice]
                    print(f"You selected: {selected_file}")
                    return selected_file
                else: print("Invalid number. Please try again.")
            except ValueError: print("Invalid input. Please enter a number.")
    except Exception as e:
        print(f"An error occurred while selecting a file: {e}")
        return None

def get_target_email_count():
    """Prompts the user for the total number of items to generate."""
    while True:
        try:
            count = int(input("Enter the total number of items (emails, chats, invites) you want to generate: "))
            if count > 0: return count
            else: print("Please enter a positive number.")
        except ValueError: print("Invalid input. Please enter a whole number.")

def get_chat_format_preference():
    """Prompts the user to select the output format for chat data."""
    print("\nSelect Chat Data Output Format:")
    print("  1: Slack Native Export (JSON) - Best for Slack simulation")
    print("  2: RSMF (Microsoft Teams) - Best for Teams simulation")
    print("  3: Webex (API/JSON format) - Best for Cisco Webex simulation")
    print("  4: All Formats (Generate everything)")
    
    while True:
        try:
            choice = int(input("Enter a number (1-4): "))
            if choice == 1: return 'slack'
            elif choice == 2: return 'teams'
            elif choice == 3: return 'webex'
            elif choice == 4: return 'all'
            else: print("Invalid number. Please enter 1-4.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def get_log_size_preference():
    """Prompts the user for the size of stress-test log files."""
    print("\n[Stress Test Configuration]")
    print("How large should the 'Large Log' attachments be (in MB)?")
    print("Note: Sizes > 500MB may cause memory errors on standard laptops.")
    while True:
        try:
            size_mb = int(input("Enter size in MB (e.g., 50): "))
            if size_mb > 0: return size_mb
            else: print("Please enter a positive number.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def get_container_preference():
    """Prompts the user regarding nested compression."""
    print("\n[Post-Processing Stress Test]")
    print("Do you want to create a Nested Container (TarGz -> Zip -> Files)?")
    print("This tests ingestion recursion depth.")
    print("  1: Yes (Create final nested archive)")
    print("  2: No (Leave as folders)")

    while True:
        try:
            choice = int(input("Enter a number (1-2): "))
            if choice == 1: return True
            elif choice == 2: return False
            else: print("Invalid number.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def get_protocol_preference():
    """Prompts the user about generating an investigation protocol document."""
    print("\n[E-Discovery Protocol Document]")
    print("Generate a search protocol document? (recommended for EAIDA testing)")
    print("This creates a document with:")
    print("  - Search terms and Boolean queries")
    print("  - Key custodian list")
    print("  - Expected hot documents")
    print("  - Date ranges and investigation scope")
    print("  1: Yes (Generate protocol document)")
    print("  2: No (Skip)")

    while True:
        try:
            choice = int(input("Enter a number (1-2): "))
            if choice == 1: return True
            elif choice == 2: return False
            else: print("Invalid number.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def get_scenario_filter_preference():
    """Prompts the user to select which investigation type(s) to generate."""
    print("\n[Investigation Type Selection]")
    print("Which investigation type do you want to generate?")
    print("\n  1: Antitrust / Price-Fixing Investigation")
    print("     - Competitor collusion scenarios")
    print("     - Coded language and coordination")
    print("     - Executive involvement")
    print("\n  2: Product Safety Fraud Investigation")
    print("     - Safety test manipulation")
    print("     - Evidence destruction")
    print("     - Regulatory fraud")
    print("\n  3: HR / Workplace Harassment Investigation")
    print("     - Harassment complaints")
    print("     - Hostile environment")
    print("     - Witness statements")
    print("\n  4: All Scenarios Mixed (Infrastructure Stress Testing)")
    print("     - All investigation types combined")
    print("     - Best for performance/volume testing only")
    print("     - Not realistic for single investigation testing")
    print("\n  5: Custom Combination")
    print("     - Select multiple investigation types to combine")

    while True:
        try:
            choice = int(input("\nEnter a number (1-5): "))
            if choice == 1:
                print("\n✓ Selected: Antitrust Investigation Only")
                return 'antitrust'
            elif choice == 2:
                print("\n✓ Selected: Safety Fraud Investigation Only")
                return 'safety_fraud'
            elif choice == 3:
                print("\n✓ Selected: HR Misconduct Investigation Only")
                return 'hr_misconduct'
            elif choice == 4:
                print("\n✓ Selected: All Scenarios Mixed (Stress Testing)")
                return 'all'
            elif choice == 5:
                print("\nCustom Combination - Select investigation types to include:")
                print("Enter types separated by commas (e.g., 1,2 for Antitrust + Safety Fraud)")
                print("  1 = Antitrust")
                print("  2 = Safety Fraud")
                print("  3 = HR Misconduct")
                custom_input = input("Enter your selection: ").strip()

                selected_filters = []
                type_map = {'1': 'antitrust', '2': 'safety_fraud', '3': 'hr_misconduct'}

                for num in custom_input.split(','):
                    num = num.strip()
                    if num in type_map:
                        selected_filters.append(type_map[num])

                if selected_filters:
                    print(f"\n✓ Selected: {', '.join(selected_filters)}")
                    return selected_filters
                else:
                    print("Invalid selection. Please try again.")
            else:
                print("Invalid number. Please enter 1-5.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def load_config(config_path):
    """Loads the selected YAML configuration file."""
    try:
        with open(config_path, 'r') as f: return yaml.safe_load(f)
    except (FileNotFoundError, yaml.YAMLError) as e:
        print(f"Error loading or parsing YAML file '{config_path}': {e}")
        return None

def filter_scenarios_by_type(scenarios, scenario_filter, include_noise=True):
    """
    Filters scenarios based on the scenario_filter setting.

    scenario_filter can be:
    - 'all' or None: Include all scenarios
    - 'antitrust': Only price-fixing scenarios (S1, S1A, S1B) + noise + privilege
    - 'safety_fraud': Only safety fraud scenarios (S2) + noise + privilege
    - 'hr_misconduct': Only HR/workplace scenarios + noise + privilege
    - A list like ['antitrust', 'safety_fraud']: Multiple scenario types + noise + privilege
    - 'antitrust_only': Signal scenarios without noise or privilege (add '_only' suffix)

    include_noise: If True, includes contextual noise scenarios (S3-S15)
    Note: S3 (privilege) is now part of noise and included in ALL investigations
    """
    if not scenario_filter or scenario_filter == 'all':
        return scenarios

    # Define scenario type mappings (signal scenarios - what you're investigating)
    scenario_mappings = {
        'antitrust': ['(S1)', '(S1A)', '(S1B)'],
        'safety_fraud': ['(S2)'],
        'hr_misconduct': ['(S_HR)'],  # Custom tag for HR scenarios
    }

    # Noise scenarios that should be included for realistic context
    # S3 (privilege) is now part of noise - it's a review task, not an investigation
    noise_prefixes = ['(S3)', '(S4)', '(S5)', '(S6)', '(S7)', '(S8)', '(S9)',
                     '(S10)', '(S11)', '(S12)', '(S13)', '(S14)', '(S15)']

    # Handle single filter or list of filters
    if isinstance(scenario_filter, str):
        # Check for '_only' suffix to exclude noise
        if scenario_filter.endswith('_only'):
            include_noise = False
            scenario_filter = scenario_filter.replace('_only', '')
        filters = [scenario_filter]
    else:
        filters = scenario_filter

    # Collect all matching scenario prefixes
    matching_prefixes = []
    for f in filters:
        if f in scenario_mappings:
            matching_prefixes.extend(scenario_mappings[f])

    # Add noise scenarios if requested (including privilege)
    if include_noise:
        matching_prefixes.extend(noise_prefixes)

    # Filter scenarios
    filtered = []
    for scenario in scenarios:
        desc = scenario.get('description', '')
        # Check if any of the matching prefixes are in the description
        if any(prefix in desc for prefix in matching_prefixes):
            filtered.append(scenario)

    return filtered

def build_context_block(profiles):
    """Dynamically builds the context string from company profiles."""
    context = "Context for a fictional simulation:\n"
    for company in profiles:
        context += f"\nCompany: {company['name']}\n"
        for person in company['personnel']: context += f"- {person['name']}, {person['title']} ({person['email']})\n"
    return context

def build_personnel_map(profiles):
    """Maps emails/names to full profiles for easy lookup."""
    personnel_map = {}
    for company in profiles:
        for person in company['personnel']:
            personnel_map[person['name']] = person
            personnel_map[person['email']] = person
    return personnel_map

def add_prompt_variation(prompt, variation_level='medium'):
    """Adds natural variation to prompts to reduce LLM repetition."""
    starters = {'low': ['Draft', 'Write', 'Compose', 'Create'],'medium': ['Draft', 'Write', 'Compose', 'Create', 'Generate', 'Produce'],'high': ['Draft', 'Write', 'Compose', 'Create', 'Generate', 'Produce', 'Put together', 'Craft']}
    fillers = {'low': ['', ''],'medium': ['', '', 'a brief', 'a concise', 'a professional'],'high': ['', '', 'a brief', 'a concise', 'a professional', 'an appropriate', 'a clear', 'a well-written']}
    for starter in ['Draft', 'Write', 'Compose', 'Create']:
        if prompt.startswith(starter):
            new_starter = random.choice(starters.get(variation_level, starters['medium']))
            filler = random.choice(fillers.get(variation_level, fillers['medium']))
            prompt = prompt.replace(starter, f"{new_starter} {filler}" if filler else new_starter, 1)
            break
    return prompt

def get_randomized_prompt(prompt_template, variables, personnel_map, run_count=1):
    """Replaces placeholders in a prompt with random choices, ensuring sender != recipient."""
    if isinstance(prompt_template, dict) and 'prompt_templates' in prompt_template:
        prompt = random.choice(prompt_template['prompt_templates'])
    else:
        prompt = prompt_template

    if run_count > 1:
        prompt = add_prompt_variation(prompt, 'high' if run_count > 3 else 'medium')

    if not variables: return prompt

    # Create a local copy so we don't modify the original config
    local_vars = variables.copy()

    # Special logic for sender/recipient to prevent self-emailing
    if '{sender}' in prompt and '{recipient}' in prompt and 'employee_pool' in local_vars:
        pool = local_vars['employee_pool']
        # Try up to 10 times to find a pair that aren't the same person
        for _ in range(10):
            sender_name, recipient_name = random.sample(pool, 2)
            
            # Resolve names to emails to check identity (handles aliases like "T. Brooks" vs "Taylor Brooks")
            sender_email = personnel_map.get(sender_name, {}).get('email')
            recipient_email = personnel_map.get(recipient_name, {}).get('email')

            # If we found emails and they are different, OR we couldn't find emails (noise names), accept it
            if sender_email and recipient_email and sender_email != recipient_email:
                break
            if not sender_email or not recipient_email:
                # If names aren't in the map (e.g. intentional typos), assume they are distinct
                break
        
        prompt = prompt.replace('{sender}', sender_name).replace('{recipient}', recipient_name)
    
    # Handle other variables
    for key, values in local_vars.items():
        if key != 'employee_pool': # We handled this manually above
            placeholder = f"{{{key}}}"
            if placeholder in prompt: 
                prompt = prompt.replace(placeholder, random.choice(values))
                
    return prompt

def get_sender_name_from_prompt(prompt, personnel_map):
    """Attempts to find a sender's name in the prompt text."""
    match = re.search(r"from\s+([A-Za-z\s\.]+)", prompt)
    if match:
        name = match.group(1).strip().replace('.', '')
        if name in personnel_map: return name
    for name in personnel_map:
        if name in prompt: return name
    return None

try:
    client = AzureOpenAI(azure_endpoint=os.getenv("AZURE_ENDPOINT"), api_key=os.getenv("AZURE_API_KEY"), api_version=os.getenv("AZURE_API_VERSION"))
    AZURE_MODEL_NAME = os.getenv("AZURE_OPENAI_MODEL")
except Exception as e:
    print(f"Error: Failed to configure AzureOpenAI client. Check .env file. Details: {e}")
    exit()

# --- Rate Limiting and Retry Logic ---

def call_llm_with_retry(func, *args, max_retries=5, **kwargs):
    """
    Wrapper to call LLM functions with exponential backoff retry logic for 429 errors.

    Args:
        func: The LLM function to call
        *args: Positional arguments for the function
        max_retries: Maximum number of retry attempts (default: 5)
        **kwargs: Keyword arguments for the function

    Returns:
        The result from the LLM function

    Raises:
        Exception: If all retries are exhausted
    """
    for attempt in range(max_retries):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            error_str = str(e)

            # Check if it's a 429 rate limit error
            if "429" in error_str or "rate_limit" in error_str.lower() or "quota" in error_str.lower():
                if attempt < max_retries - 1:
                    # Exponential backoff: 2^attempt seconds (2, 4, 8, 16, 32)
                    wait_time = 2 ** (attempt + 1)
                    print(f"    [Rate Limit] 429 error detected. Retrying in {wait_time}s (attempt {attempt + 1}/{max_retries})...")
                    time.sleep(wait_time)
                else:
                    print(f"    [Rate Limit] Max retries ({max_retries}) reached. Giving up.")
                    raise
            else:
                # For non-429 errors, raise immediately
                raise

    # Should never reach here, but just in case
    raise Exception(f"Failed after {max_retries} retries")

# --- Attachment Generation Functions ---

def generate_attachment_text_from_llm(filename, description, file_type, email_context=None, temperature=0.8):
    """Generates realistic text content for a document based on its description and email context."""
    system_message = (
    f"You are an AI assistant generating the internal text content for a fake file. "
    f"The filename is '{filename}'. "
    "Do not write an email. Do not include markdown formatting. "
    "Write realistic content appropriate for the file type. "
    "If it is a medical record, use clinical terminology and a structured format (Diagnosis, Findings, Plan). "
    "If it is a business doc, use professional corporate language."
)

    # Enhanced prompt with email context for alignment
    if email_context:
        prompt = f"""The email this attachment belongs to says:

"{email_context['body'][:800]}"

Subject: {email_context.get('subject', 'No subject')}

Generate content for: {description}

CRITICAL ALIGNMENT REQUIREMENT:
- The attachment content MUST match what the email describes
- If the email mentions specific data, findings, or conclusions, include them in the attachment
- If the email says "see attached report showing catastrophic failure", the attachment must show that failure data
- If the email references specific numbers, dates, or metrics, include those exact details
- Make the attachment tell the same story as the email, but in document format"""
    else:
        prompt = f"Generate content for this document: {description}"

    print(f"  ... Generating content for attachment: {filename} ...")
    try:
        def _call_api():
            return client.chat.completions.create(
                model=AZURE_MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                temperature=temperature
            )

        response = call_llm_with_retry(_call_api)
        return response.choices[0].message.content
    except Exception as e:
        print(f"!!! ERROR generating attachment content: {e}")
        return description # Fallback

def create_fake_pdf_attachment(filename, content_text):
    """Creates a valid PDF with basic manual text wrapping and character sanitization."""
    
    # --- SANITIZATION START ---
    replacements = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '--', # Em dash
        '\u2026': '...',# Ellipsis
        '\u00A0': ' '   # Non-breaking space
    }
    for char, replacement in replacements.items():
        content_text = content_text.replace(char, replacement)
    
    clean_text = content_text.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
    # --- SANITIZATION END ---
    
    words = clean_text.split()
    lines = []
    current_line = []
    current_length = 0
    
    for word in words:
        if current_length + len(word) > 85: # Approx char limit per line
            lines.append(" ".join(current_line))
            current_line = [word]
            current_length = len(word)
        else:
            current_line.append(word)
            current_length += len(word) + 1
    lines.append(" ".join(current_line))
    
    pdf_text_stream = "BT /F1 11 Tf 50 750 Td 15 TL " 
    for line in lines:
        pdf_text_stream += f"({line}) Tj T* "
    pdf_text_stream += "ET"

    pdf_content = f"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R>>endobj
4 0 obj<</Length {len(pdf_text_stream) + 20}>>stream
{pdf_text_stream}
endstream endobj
xref
0 5
0000000000 65535 f
0000000010 00000 n
0000000059 00000 n
0000000112 00000 n
0000000199 00000 n
trailer<</Size 5/Root 1 0 R>>
startxref
288
%%EOF"""
    
    return pdf_content.encode('latin-1', errors='replace')

def create_fake_word_doc(filename, content_text):
    """Creates a fake .docx file in memory with generated content."""
    document = Document()
    document.add_heading(filename, level=1)
    
    # Add paragraphs
    for para in content_text.split('\n'):
        if para.strip():
            document.add_paragraph(para.strip())
            
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def create_fake_excel_sheet(filename, content_text):
    """Creates a fake .xlsx file with generated content and dummy data."""
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report Data"
    sheet['A1'] = filename
    
    # Put the text content into rows at the top
    current_row = 3
    for line in content_text.split('\n'):
        if line.strip():
            sheet[f'A{current_row}'] = line
            current_row += 1
            
    current_row += 2
    headers = ['ID', 'Category', 'Amount', 'Status', 'Review Date']
    for col, header in enumerate(headers, start=1):
        sheet.cell(row=current_row, column=col, value=header)
        
    # Add some random rows of "data"
    for i in range(1, 20):
        row_num = current_row + i
        sheet.cell(row=row_num, column=1, value=f"REC-{1000+i}")
        sheet.cell(row=row_num, column=2, value=random.choice(['Hardware', 'Software', 'Services', 'Logistics']))
        sheet.cell(row=row_num, column=3, value=random.randint(500, 50000))
        sheet.cell(row=row_num, column=4, value=random.choice(['Approved', 'Pending', 'Rejected']))
        sheet.cell(row=row_num, column=5, value=datetime.now().strftime("%Y-%m-%d"))
        
    file_stream = BytesIO()
    workbook.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def create_near_duplicate(email_content, variation_type='signature'):
    """Creates near-duplicate by modifying non-substantive content."""
    modified_content = email_content.copy()
    variation_options = ['signature', 'disclaimer', 'formatting']
    variation_type = random.choice(variation_options)
    if variation_type == 'signature':
        if "\n\nSent from my iPhone" not in modified_content['body']:
             modified_content['body'] += "\n\nSent from my iPhone"
        else:
             modified_content['body'] = modified_content['body'].replace("\n\nSent from my iPhone", "")
    elif variation_type == 'disclaimer':
        disclaimer = "\n\n---\nThis email and any attachments are confidential and intended solely for the use of the individual or entity to whom they are addressed."
        if disclaimer not in modified_content['body']:
            modified_content['body'] += disclaimer
    elif variation_type == 'formatting':
        if '\n\n' in modified_content['body']:
            modified_content['body'] = modified_content['body'].replace('\n\n', '\n')
        else:
            modified_content['body'] = modified_content['body'].replace('\n', '\n\n')
    return modified_content

def inject_blast_recipients(email_content, personnel_map):
    """
    STRESS TEST FEATURE:
    Programmatically appends up to 500 email addresses to the recipients list.
    """
    print("  -> Stress Test Triggered: Injecting 500+ recipients...")
    # Extract all emails from the personnel map
    all_emails = [p['email'] for p in personnel_map.values() if 'email' in p]
    
    # If we have fewer than 500, just repeat them to reach 500
    while len(all_emails) < 500:
        all_emails.extend(all_emails)
    
    target_list = all_emails[:500]
    
    # Append to existing recipients
    current_recipients = email_content.get('recipients', [])
    for email in target_list:
        # Use simple names if map lookup fails
        name = personnel_map.get(email, {}).get('name', 'Employee')
        current_recipients.append([name, email])
        
    email_content['recipients'] = current_recipients
    return email_content

# --- LLM and Content Generation ---

def format_quoted_body(previous_content, previous_date):
    """Formats the previous email's content into a standard reply quote."""
    sender_line = f"From: {previous_content['sender_name']} <{previous_content['sender_email']}>"
    date_line = f"Sent: {previous_date.strftime('%A, %B %d, %Y %I:%M %p')}"
    to_line = f"To: {', '.join([f'{name} <{email}>' for name, email in previous_content['recipients']])}"
    
    cc_line = ""
    if 'cc_recipients' in previous_content and previous_content['cc_recipients']:
        cc_list = ', '.join([f'{name} <{email}>' for name, email in previous_content['cc_recipients']])
        cc_line = f"\nCc: {cc_list}"

    subject_line = f"Subject: {previous_content['subject']}"
    quoted_body = '\n'.join([f"> {line}" for line in previous_content['body'].splitlines()])
    
    return f"\n\n-----Original Message-----\n{sender_line}\n{date_line}\n{to_line}{cc_line}\n{subject_line}\n\n{quoted_body}"

def get_temperature_for_scenario(scenario_type, is_noise=False, config_temp=None):
    """Returns appropriate temperature based on scenario characteristics.

    Args:
        scenario_type: Type of scenario ('thread', 'chat', 'standalone', 'calendar')
        is_noise: Whether this is a noise scenario
        config_temp: Temperature override from YAML llm_settings (takes precedence)

    Returns:
        float: Temperature value for LLM generation
    """
    # Config override takes precedence
    if config_temp is not None:
        return config_temp

    # Default behavior (backwards compatible)
    if is_noise: return random.uniform(0.9, 1.3)
    elif scenario_type == 'thread': return 0.85
    elif scenario_type == 'chat': return 0.7  # High temp for chat spontaneity
    else: return 0.95

def generate_llm_response(prompt, system_message, temperature=0.95):
    """Generic function to get a JSON response from the LLM."""
    print(f"---> Sending prompt to Azure OpenAI Model (temp={temperature:.2f})...")
    try:
        def _call_api():
            return client.chat.completions.create(
                model=AZURE_MODEL_NAME,
                messages=[{"role": "system", "content": system_message}, {"role": "user", "content": prompt}],
                temperature=temperature,
                response_format={"type": "json_object"}
            )

        response = call_llm_with_retry(_call_api)
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"!!! ERROR: Failed to get valid response from LLM. Details: {e}")
        return None

def generate_email_content_from_llm(prompt, temperature=0.95):
    """Generates email content from LLM."""
    system_message = "You are an AI assistant for generating simulated corporate emails for a fictional story. Return a single, valid JSON object and nothing else. The JSON object must have the keys: 'subject', 'body', 'sender_name', 'sender_email', 'recipients'. 'recipients' must be a list of lists, like [['Recipient Name', 'recipient@email.com']]. You can OPTIONALLY include 'cc_recipients' and 'bcc_recipients' keys, following the same format as 'recipients'. When asked to reply, your 'body' should ONLY contain the new reply content."
    email_data = generate_llm_response(prompt, system_message, temperature)
    if email_data and all(key in email_data for key in ["subject", "body", "sender_name", "sender_email", "recipients"]):
        return email_data
    else:
        print("!!! ERROR: LLM response for email was missing required keys.")
        return None

def generate_calendar_content_from_llm(prompt, temperature=0.9):
    """Generates calendar event content from LLM."""
    system_message = "You are an AI assistant for generating simulated corporate calendar events for a fictional story. Return a single, valid JSON object and nothing else. The JSON object must have the keys: 'summary' (the event title), 'description' (event details), 'organizer_name', 'organizer_email', and 'attendees'. 'attendees' must be a list of lists, like [['Attendee Name', 'attendee@email.com']]."
    event_data = generate_llm_response(prompt, system_message, temperature)
    if event_data and all(key in event_data for key in ["summary", "description", "organizer_name", "organizer_email", "attendees"]):
        return event_data
    else:
        print("!!! ERROR: LLM response for calendar event was missing required keys.")
        return None

def generate_chat_content_from_llm(prompt, temperature=0.7):
    """Generates a back-and-forth chat conversation with realistic chat patterns."""
    system_message = (
        "You are an AI assistant for generating simulated corporate chat logs (Slack/Teams). "
        "Return a single, valid JSON object and nothing else. "
        "The JSON object must have a key 'messages', which is a list of objects. "
        "Each object in the list must have: 'sender_name', 'sender_email', 'body' (the text), and optional 'thread_ts' (for threaded replies). "

        "CRITICAL CHAT REALISM REQUIREMENTS:\n"
        "1. SHORT MESSAGES: Most messages should be 1-3 sentences MAX. Simulate rapid-fire texting style, not email paragraphs.\n"
        "2. BREAK UP THOUGHTS: If someone has multiple points, break into SEPARATE messages sent seconds apart:\n"
        "   Example: Message 1: 'hey can you look at this?', Message 2: 'the test results are concerning', Message 3: 'we should discuss offline'\n"
        "3. NATURAL INTERRUPTIONS: People reply before others finish their thought. Don't wait for complete paragraphs.\n"
        "4. CASUAL TONE: Use lowercase, abbreviations (btw, fyi, lol), occasional typos, emoji (👍 😅 🔥 etc)\n"
        "5. QUICK RESPONSES: 'ok', 'sounds good', 'got it', 'will do', 'on it' are common\n"
        "6. THREAD_TS: For replies within a conversation, set 'thread_ts' to the timestamp of the parent message being replied to\n"
        "7. REALISTIC PATTERNS:\n"
        "   - Questions get quick 'yep' / 'no' / 'checking' responses\n"
        "   - Links/attachments get 'thanks' / '👍' reactions\n"
        "   - Urgent topics have rapid back-and-forth\n"
        "   - Typing multiple messages in a row is NORMAL in chat\n"

        "BAD (email-style): 'Hi Jamie, I wanted to follow up on our discussion about the pricing strategy. I think we should consider...'\n"
        "GOOD (chat-style): 'hey jamie', 'following up on pricing', 'think we should revisit that 15% number', 'thoughts?'"
    )
    chat_data = generate_llm_response(prompt, system_message, temperature)
    if chat_data and 'messages' in chat_data and isinstance(chat_data['messages'], list):
        return chat_data
    else:
        print("!!! ERROR: LLM response for chat was missing required keys.")
        return None

def add_realistic_email_metadata(msg, scenario_description):
    """Adds realistic email client headers and importance flags."""
    msg['X-Mailer'] = random.choice(['Microsoft Outlook 16.0', 'Apple Mail (16.0.3)', 'Mozilla Thunderbird 115.3.1', 'Google Mail'])
    msg['X-Priority'] = '1 (Highest)' if random.random() < 0.1 else '3 (Normal)'
    msg['MIME-Version'] = '1.0'
    hot_keywords = ['confidential', 'privileged', 'fraud', 'urgent', 'legal', 'price-fixing', 'safety']
    if any(keyword in scenario_description.lower() for keyword in hot_keywords):
        msg['Importance'] = 'High'
        if 'legal' in scenario_description.lower() or 'privilege' in scenario_description.lower():
            msg['Sensitivity'] = 'Private'

# --- File Creation and Saving Logic ---

def create_and_save_email(base_filename, email_content, output_dir, email_date, personnel_map, scenario_description, attachment_config=None, headers=None, stats=None):
    """Creates an email, saves it, adds attachments, and updates stats."""
    sender_profile = personnel_map.get(email_content.get('sender_email'))
    if sender_profile and sender_profile.get('signature') and random.random() < 0.6:
        email_content['body'] += f"\n\n-- \n{sender_profile['signature']}"
    
    msg = EmailMessage()
    msg['Subject'] = email_content.get('subject', 'No Subject')
    msg['From'] = f"{email_content.get('sender_name', 'Unknown Sender')} <{email_content.get('sender_email', 'unknown@sender.com')}>"
    msg['To'] = ", ".join([f"{name} <{email}>" for name, email in email_content.get('recipients', [])])
    if 'cc_recipients' in email_content and email_content['cc_recipients']:
        msg['Cc'] = ", ".join([f"{name} <{email}>" for name, email in email_content['cc_recipients']])
    if 'bcc_recipients' in email_content and email_content['bcc_recipients']:
        msg['Bcc'] = ", ".join([f"{name} <{email}>" for name, email in email_content['bcc_recipients']])
    msg['Date'] = email_date.strftime("%a, %d %b %Y %H:%M:%S %z")
    
    if headers:
        for key, value in headers.items(): msg[key] = value
    
    add_realistic_email_metadata(msg, scenario_description)
    msg.set_content(email_content.get('body', ''), cte='quoted-printable')

    if stats:
        stats['emails'] += 1
        # Track date for date range calculation
        stats['email_dates'].append(email_date)
        # Track custodian (sender email)
        stats['custodians'].add(email_content.get('sender_email'))

    # --- Attachment Logic ---
    if attachment_config:
        # Retrieve the user-configured log size, default to 50MB if not set
        log_size_mb = attachment_config.get('log_size_mb', 50)
        
        for att_type in attachment_config.get('types', []):
            allowed_scenarios = att_type.get('limit_to_scenarios', [])
            if scenario_description not in allowed_scenarios:
                continue
            
            if random.random() < att_type.get('probability', 1.0):
                att_filename = random.choice(att_type['filenames']).format(
                    date=email_date.strftime('%Y-%m-%d'),
                    quarter= (email_date.month - 1) // 3 + 1,
                    version=random.randint(1, 5)
                )
                
                mime_type = att_type.get('mime_type', 'application/octet-stream')
                
                # --- STRESS TEST: LARGE LOG FILE GENERATION ---
                if mime_type == 'text/x-log':
                    # Check if YAML specifies a fixed size for this specific file type
                    # If yes, use it. If no, fall back to the user's input (log_size_mb)
                    current_file_size_mb = att_type.get('fixed_size_mb', log_size_mb)
                    
                    print(f"  -> Generating LOG file: {att_filename} ({current_file_size_mb} MB)...")
                    
                    # 1 MB = 1024 * 1024 bytes
                    target_size_bytes = current_file_size_mb * 1024 * 1024
                    
                    # Create a base string (~100 bytes)
                    base_line = f"[{datetime.now().isoformat()}] INFO: System heartbeat check - Status OK - Process ID {random.randint(1000,9999)} - Memory: {random.randint(200,800)}MB\n"
                    base_bytes = base_line.encode('utf-8')
                    
                    # Calculate how many iterations needed to reach target MB
                    iterations = int(target_size_bytes / len(base_bytes))
                    
                    # Generate content in one mathematical operation (fast)
                    full_content = base_bytes * iterations
                    
                    msg.add_attachment(full_content, maintype='text', subtype='plain', filename=att_filename)
                    if stats:
                        stats['attachments'] += 1
                        stats['attachment_types']['.log'] = stats['attachment_types'].get('.log', 0) + 1
                    
                    continue # Skip LLM generation and standard processing
                # ---------------------------------------------

                short_desc = att_type.get('content_description', f'Content for {att_filename}')

                # Generate realistic content with email context for alignment
                file_context = "Word Document"
                if "spreadsheet" in mime_type: file_context = "Excel Spreadsheet Data"
                if "pdf" in mime_type: file_context = "PDF Document"

                # Pass email content to ensure attachment aligns with email narrative
                full_content_text = generate_attachment_text_from_llm(
                    att_filename,
                    short_desc,
                    file_context,
                    email_context=email_content  # NEW: Pass email context for alignment
                )
                
                file_data = None
                subtype = None
                ext = 'unknown'

                if mime_type == 'application/pdf':
                    print(f"  -> Attaching file: {att_filename}")
                    file_data = create_fake_pdf_attachment(att_filename, full_content_text)
                    subtype = 'pdf'
                    ext = '.pdf'
                
                elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    print(f"  -> Attaching file: {att_filename}")
                    file_data = create_fake_word_doc(att_filename, full_content_text)
                    subtype = 'vnd.openxmlformats-officedocument.wordprocessingml.document'
                    ext = '.docx'
                
                elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                    print(f"  -> Attaching file: {att_filename}")
                    file_data = create_fake_excel_sheet(att_filename, full_content_text)
                    subtype = 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    ext = '.xlsx'

                if file_data:
                    msg.add_attachment(file_data, maintype='application', subtype=subtype, filename=att_filename)
                    if stats:
                        stats['attachments'] += 1
                        stats['attachment_types'][ext] = stats['attachment_types'].get(ext, 0) + 1
    
    # Exact Duplicate (Custodian Folder) Logic
    all_custodians = []

    # Add sender if valid
    sender_email = email_content.get('sender_email')
    if sender_email and isinstance(sender_email, str) and '@' in sender_email:
        all_custodians.append(sender_email)

    # Add recipients with validation
    for recipient_pair in email_content.get('recipients', []):
        if isinstance(recipient_pair, (list, tuple)) and len(recipient_pair) >= 2:
            email = recipient_pair[1]
            if email and isinstance(email, str) and '@' in email:
                all_custodians.append(email)

    # Add CC recipients with validation
    if 'cc_recipients' in email_content:
        for cc_pair in email_content.get('cc_recipients', []):
            if isinstance(cc_pair, (list, tuple)) and len(cc_pair) >= 2:
                email = cc_pair[1]
                if email and isinstance(email, str) and '@' in email:
                    all_custodians.append(email)

    # Add BCC recipients with validation
    if 'bcc_recipients' in email_content:
        for bcc_pair in email_content.get('bcc_recipients', []):
            if isinstance(bcc_pair, (list, tuple)) and len(bcc_pair) >= 2:
                email = bcc_pair[1]
                if email and isinstance(email, str) and '@' in email:
                    all_custodians.append(email)

    email_as_string = str(msg)

    # Create folders only for valid, unique email addresses
    for email_address in set(all_custodians):
        custodian_folder_name = email_address.split('@')[0]
        custodian_path = os.path.join(output_dir, custodian_folder_name)
        os.makedirs(custodian_path, exist_ok=True)
        filepath = os.path.join(custodian_path, f"{base_filename}.eml")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(email_as_string)
                
    return msg.get('Message-ID')

def create_and_save_calendar_event(filename, event_content, output_dir, event_date, stats=None):
    """Creates and saves a calendar event as an .ics file."""
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    if not filename.endswith('.ics'): filename = f"{filename}.ics"
    dtstamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    dtstart = event_date.strftime("%Y%m%dT%H%M%SZ")
    dtend = (event_date + timedelta(hours=1)).strftime("%Y%m%dT%H%M%SZ")
    attendee_lines = [f"ATTENDEE;CN={name};ROLE=REQ-PARTICIPANT:mailto:{email}" for name, email in event_content.get('attendees', [])]
    ics_content = ["BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//MySyntheticDataGenerator//EN", "BEGIN:VEVENT", f"UID:{uuid.uuid4()}@mygenerator.com", f"DTSTAMP:{dtstamp}", f"ORGANIZER;CN={event_content.get('organizer_name', 'Unknown')}:mailto:{event_content.get('organizer_email', 'unknown@organizer.com')}", *attendee_lines, f"DTSTART:{dtstart}", f"DTEND:{dtend}", f"SUMMARY:{event_content.get('summary', 'No Summary')}", f"DESCRIPTION:{event_content.get('description', '').replace(chr(10), chr(92)+'n')}", "END:VEVENT", "END:VCALENDAR"]
    full_path = os.path.join(output_dir, filename)
    with open(full_path, 'w', encoding='utf-8') as f: f.write("\n".join(ics_content))
    
    if stats:
        stats['calendar_events'] += 1
        # Track date for date range calculation
        stats['email_dates'].append(event_date)
        # Track custodian (organizer email)
        stats['custodians'].add(event_content.get('organizer_email'))

# --- Chat Generation Functions ---

def get_deterministic_id(key, prefix="U"):
    """Generates a consistent 9-char ID based on an email or name string."""
    seed = sum(ord(c) for c in key)
    random.seed(seed)
    suffix = ''.join(random.choices('0123456789ABCDEF', k=8))
    random.seed() 
    return f"{prefix}{suffix}"

def create_and_save_slack_native(base_filename, chat_content, output_dir, start_date, personnel_map, stats=None):
    """
    Creates a Native Slack Export folder structure with modern Block Kit formatting.
    Structure: /slack_export/users.json, channels.json, /channel/YYYY-MM-DD.json
    """
    slack_root = os.path.join(output_dir, "slack_export")
    os.makedirs(slack_root, exist_ok=True)

    # Generate a consistent Team ID
    random.seed(output_dir)
    team_suffix = ''.join(random.choices('0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ', k=10))
    team_id = f"T{team_suffix}"
    random.seed() # Reset seed

    # 1. Generate users.json
    users_list = []
    for email, profile in personnel_map.items():
        user_id = get_deterministic_id(email, "U")
        first_name = profile['name'].split()[0]
        real_name = profile['name']
        display_name = first_name.lower()
        avatar_hash = "".join(random.choices('0123456789abcdef', k=12))
        
        users_list.append({
            "id": user_id, "team_id": team_id, "name": email.split('@')[0], "deleted": False, "color": "9f69e7", "real_name": real_name, "tz": "America/Los_Angeles", "tz_label": "Pacific Daylight Time", "tz_offset": -25200,
            "profile": {
                "title": profile.get('title', 'Employee'), "phone": "", "skype": "", "real_name": real_name, "real_name_normalized": real_name, "display_name": display_name, "display_name_normalized": display_name, "fields": {}, "status_text": "", "status_emoji": "", "status_expiration": 0, "avatar_hash": avatar_hash, "image_24": "https://secure.gravatar.com/avatar/example.jpg?s=24", "image_32": "https://secure.gravatar.com/avatar/example.jpg?s=32", "image_48": "https://secure.gravatar.com/avatar/example.jpg?s=48", "image_72": "https://secure.gravatar.com/avatar/example.jpg?s=72", "image_192": "https://secure.gravatar.com/avatar/example.jpg?s=192", "image_512": "https://secure.gravatar.com/avatar/example.jpg?s=512", "status_text_canonical": "", "team": team_id, "first_name": first_name, "last_name": profile['name'].split()[-1] if ' ' in profile['name'] else "", "is_custom_image": False, "email": email
            },
            "is_admin": False, "is_owner": False, "is_primary_owner": False, "is_restricted": False, "is_ultra_restricted": False, "is_bot": False, "is_app_user": False, "updated": int(start_date.timestamp()), "is_email_confirmed": True
        })
    
    # Merge logic for users.json
    users_file_path = os.path.join(slack_root, "users.json")
    if os.path.exists(users_file_path):
        try:
            with open(users_file_path, 'r', encoding='utf-8') as f:
                existing_users = json.load(f)
                existing_ids = {u['id'] for u in existing_users}
                for new_user in users_list:
                    if new_user['id'] not in existing_ids: existing_users.append(new_user)
                users_list = existing_users
        except: pass

    with open(users_file_path, 'w', encoding='utf-8') as f: json.dump(users_list, f, indent=4)

    # 2. Determine Channel
    channel_name = re.sub(r'[^a-z0-9-_]', '', base_filename.lower().replace(' ', '-'))[:21]
    channel_id = get_deterministic_id(channel_name, "C")
    
    # Update channels.json
    channels_file = os.path.join(slack_root, "channels.json")
    channels_list = []
    if os.path.exists(channels_file):
        try:
            with open(channels_file, 'r') as f: channels_list = json.load(f)
        except: pass
    
    if not any(c['id'] == channel_id for c in channels_list):
        creator_id = users_list[0]['id'] if users_list else "U00000000"
        channels_list.append({
            "id": channel_id, "name": channel_name, "is_channel": True, "created": int(start_date.timestamp()), "is_archived": False, "is_general": False, "unlinked": 0, "creator": creator_id, "is_shared": False, "is_org_shared": False, "is_member": True, "is_private": False, "is_mpim": False, "members": [u['id'] for u in users_list],
            "topic": {"value": f"Topic for {channel_name}", "creator": creator_id, "last_set": int(start_date.timestamp())},
            "purpose": {"value": f"Purpose of {channel_name}", "creator": creator_id, "last_set": int(start_date.timestamp())}
        })
        with open(channels_file, 'w', encoding='utf-8') as f: json.dump(channels_list, f, indent=4)

    # 3. Create Messages with Enhanced Realism
    messages = []
    current_time = start_date
    message_timestamps = {}  # Track timestamps for threading

    for idx, msg in enumerate(chat_content.get('messages', [])):
        # Realistic chat timing: rapid-fire (5-30 sec) or thoughtful pauses (1-5 min)
        msg_length = len(msg.get('body', ''))
        if msg_length < 30:  # Short messages = rapid fire
            current_time += timedelta(seconds=random.randint(5, 30))
        elif msg_length < 100:  # Medium messages
            current_time += timedelta(seconds=random.randint(20, 120))
        else:  # Longer messages = more thinking time
            current_time += timedelta(seconds=random.randint(60, 300))

        ts_val = f"{current_time.timestamp():.6f}"
        sender_email = msg.get('sender_email', 'unknown@chat.com')
        user_id = get_deterministic_id(sender_email, "U")
        real_name = msg.get('sender_name', 'Unknown User')
        first_name = real_name.split()[0]
        msg_text = msg.get('body', '')
        client_msg_id = str(uuid.uuid4())

        user_profile_data = {
            "avatar_hash": "".join(random.choices('0123456789abcdef', k=12)), "image_72": "https://secure.gravatar.com/avatar/example.jpg?s=72", "first_name": first_name, "real_name": real_name, "display_name": first_name.lower(), "team": team_id, "name": sender_email.split('@')[0], "is_restricted": False, "is_ultra_restricted": False
        }
        blocks = [{"type": "rich_text", "block_id": "".join(random.choices('0123456789abcdefghijklmnopqrstuvwxyz', k=5)), "elements": [{"type": "rich_text_section", "elements": [{"type": "text", "text": msg_text}]}]}]

        slack_msg = {
            "client_msg_id": client_msg_id, "type": "message", "text": msg_text, "user": user_id, "ts": ts_val, "team": team_id, "user_team": team_id, "source_team": team_id, "user_profile": user_profile_data, "blocks": blocks
        }

        # Handle threading (if thread_ts provided by LLM)
        thread_ts = msg.get('thread_ts')
        if thread_ts and thread_ts in message_timestamps:
            slack_msg["thread_ts"] = message_timestamps[thread_ts]
            slack_msg["parent_user_id"] = message_timestamps.get(f"{thread_ts}_user")

        # Store timestamp for potential threading
        message_timestamps[str(idx)] = ts_val
        message_timestamps[f"{idx}_user"] = user_id

        # Add reactions (30% chance for short confirmations, 15% for other messages)
        if random.random() < (0.3 if msg_length < 20 else 0.15):
            reaction_emoji = random.choice([
                'thumbsup', '+1', 'white_check_mark', 'ok_hand', 'fire',
                'eyes', 'point_up', 'raised_hands', 'slightly_smiling_face', 'sweat_smile'
            ])
            # Pick 1-3 random users to react
            num_reactors = random.randint(1, min(3, len(users_list)))
            reactors = random.sample([u['id'] for u in users_list], num_reactors)
            slack_msg["reactions"] = [{
                "name": reaction_emoji,
                "users": reactors,
                "count": len(reactors)
            }]

        # Add edit history (10% chance for messages over 50 chars)
        if random.random() < 0.1 and msg_length > 50:
            edit_time = current_time + timedelta(seconds=random.randint(30, 300))
            slack_msg["edited"] = {
                "user": user_id,
                "ts": f"{edit_time.timestamp():.6f}"
            }

        # Handle special subtypes
        if "joined" in msg_text.lower() and len(msg_text) < 50:
             slack_msg["subtype"] = "channel_join"
             slack_msg.pop("blocks", None)
             slack_msg.pop("client_msg_id", None)

        messages.append(slack_msg)

    # 4. Save Daily Log
    channel_dir = os.path.join(slack_root, channel_name)
    os.makedirs(channel_dir, exist_ok=True)
    date_filename = start_date.strftime("%Y-%m-%d.json")
    final_path = os.path.join(channel_dir, date_filename)
    
    existing_msgs = []
    if os.path.exists(final_path):
        try:
            with open(final_path, 'r', encoding='utf-8') as f: existing_msgs = json.load(f)
        except: pass
    
    existing_msgs.extend(messages)
    existing_msgs.sort(key=lambda x: float(x['ts']))

    with open(final_path, 'w', encoding='utf-8') as f: json.dump(existing_msgs, f, indent=4)
        
    if stats:
        stats['scenarios_triggered'][f"Slack: {channel_name}"] = stats['scenarios_triggered'].get(f"Slack: {channel_name}", 0) + 1
        stats['rsmf_chats'] = stats.get('rsmf_chats', 0) + 1
        # Track date for date range calculation
        stats['email_dates'].append(start_date)
        # Track custodians (participants in chat)
        for msg in chat_content.get('messages', []):
            stats['custodians'].add(msg.get('sender_email'))

    print(f"  -> Created Native Slack Export: {channel_name}/{date_filename} (Modern Format)")

def create_and_save_rsmf(base_filename, chat_content, output_dir, start_date, personnel_map, stats=None):
    """Creates a Relativity Short Message Format (RSMF) file."""
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    
    conversation_id = str(uuid.uuid4())
    participants = []
    events = []
    participant_map = {}
    
    unique_senders = {}
    for msg in chat_content.get('messages', []):
        email = msg.get('sender_email', 'unknown@chat.com')
        name = msg.get('sender_name', 'Unknown')
        if email not in unique_senders: unique_senders[email] = name

    for email, name in unique_senders.items():
        p_id = email
        participant_map[email] = p_id
        participants.append({"id": p_id, "display": name, "email": email, "type": "User"})

    current_time = start_date
    for msg in chat_content.get('messages', []):
        # Realistic chat timing based on message length
        msg_length = len(msg.get('body', ''))
        if msg_length < 30:
            current_time += timedelta(seconds=random.randint(5, 30))
        elif msg_length < 100:
            current_time += timedelta(seconds=random.randint(20, 120))
        else:
            current_time += timedelta(seconds=random.randint(60, 300))

        iso_timestamp = current_time.strftime("%Y-%m-%dT%H:%M:%S%z")
        if not iso_timestamp.endswith('Z') and '+' not in iso_timestamp and '-' not in iso_timestamp: iso_timestamp += "Z"

        event = {
            "type": "message", "id": str(uuid.uuid4()), "participant": participant_map.get(msg.get('sender_email')), "body": msg.get('body'), "timestamp": iso_timestamp, "conversation": conversation_id
        }

        # Add reactions for Teams (similar to Slack)
        if random.random() < (0.3 if msg_length < 20 else 0.15):
            reaction_type = random.choice(['like', 'heart', 'laugh', 'surprised', 'sad'])
            num_reactors = random.randint(1, min(3, len(participants)))
            reactors = random.sample([p['id'] for p in participants], num_reactors)
            event["reactions"] = [{
                "type": reaction_type,
                "users": reactors
            }]

        # Add edited flag (10% chance for longer messages)
        if random.random() < 0.1 and msg_length > 50:
            edit_time = current_time + timedelta(seconds=random.randint(30, 300))
            event["editedTimestamp"] = edit_time.strftime("%Y-%m-%dT%H:%M:%S%z")

        events.append(event)

    manifest = {
        "version": "1.0.0",
        "conversations": [{"id": conversation_id, "display": f"Chat - {base_filename}", "platform": "Microsoft Teams", "type": "Direct", "participants": [p['id'] for p in participants]}],
        "participants": participants, "events": events
    }

    custodian_email = participants[0]['email'] if participants else "unknown"
    custodian_folder = custodian_email.split('@')[0]
    custodian_path = os.path.join(output_dir, custodian_folder)
    os.makedirs(custodian_path, exist_ok=True)

    rsmf_filename = f"{base_filename}.rsmf"
    full_path = os.path.join(custodian_path, rsmf_filename)

    try:
        with zipfile.ZipFile(full_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('rsmf_manifest.json', json.dumps(manifest, indent=4))
        
        if stats:
            stats['rsmf_chats'] = stats.get('rsmf_chats', 0) + 1
            # Track date for date range calculation
            stats['email_dates'].append(start_date)
            # Track custodians (participants in chat)
            for participant in participants:
                stats['custodians'].add(participant.get('email'))

        print(f"  -> Created RSMF Chat log: {rsmf_filename}")
    except Exception as e:
        print(f"!!! Error creating RSMF file: {e}")

def create_and_save_webex_native(base_filename, chat_content, output_dir, start_date, personnel_map, stats=None):
    """
    Creates a Simulated Webex API Export structure.
    Structure: /webex_export/rooms.json, participants.json, messages.json
    """
    webex_root = os.path.join(output_dir, "webex_export", base_filename)
    os.makedirs(webex_root, exist_ok=True)

    room_id = get_deterministic_id(base_filename, "Y2lzY")
    room_title = base_filename.replace('_', ' ').title()
    
    rooms_data = [{
        "id": room_id, "title": room_title, "type": "group", "created": start_date.strftime("%Y-%m-%dT%H:%M:%S.000Z"), "lastActivity": (start_date + timedelta(hours=1)).strftime("%Y-%m-%dT%H:%M:%S.000Z"), "isArchived": False
    }]

    unique_senders = {}
    for msg in chat_content.get('messages', []):
        email = msg.get('sender_email', 'unknown@webex.com')
        name = msg.get('sender_name', 'Unknown')
        if email not in unique_senders: unique_senders[email] = name

    participants_data = []
    for email, name in unique_senders.items():
        p_id = get_deterministic_id(email, "Y2lzY")
        participants_data.append({
            "id": p_id, "displayName": name, "emails": [email], "avatar": None, "created": (start_date - timedelta(days=30)).strftime("%Y-%m-%dT%H:%M:%S.000Z")
        })

    messages_data = []
    current_time = start_date
    for msg in chat_content.get('messages', []):
        current_time += timedelta(seconds=random.randint(5, 120))
        iso_time = current_time.strftime("%Y-%m-%dT%H:%M:%S.000Z")
        sender_email = msg.get('sender_email', 'unknown@webex.com')
        person_id = get_deterministic_id(sender_email, "Y2lzY")
        msg_id = str(uuid.uuid4())

        webex_msg = {
            "id": msg_id, "roomId": room_id, "roomType": "group", "text": msg.get('body', ''), "personId": person_id, "personEmail": sender_email, "created": iso_time, "edited": None, "deleted": False, "parentId": None, "attachments": [] 
        }
        messages_data.append(webex_msg)

    with open(os.path.join(webex_root, "rooms.json"), 'w', encoding='utf-8') as f: json.dump(rooms_data, f, indent=4)
    with open(os.path.join(webex_root, "participants.json"), 'w', encoding='utf-8') as f: json.dump(participants_data, f, indent=4)
    with open(os.path.join(webex_root, "messages.json"), 'w', encoding='utf-8') as f: json.dump(messages_data, f, indent=4)

    if stats:
        stats['scenarios_triggered'][f"Webex: {room_title}"] = stats['scenarios_triggered'].get(f"Webex: {room_title}", 0) + 1
        stats['rsmf_chats'] = stats.get('rsmf_chats', 0) + 1
        # Track date for date range calculation
        stats['email_dates'].append(start_date)
        # Track custodians (participants in chat)
        for msg in chat_content.get('messages', []):
            stats['custodians'].add(msg.get('sender_email'))

    print(f"  -> Created Webex API Export: {base_filename}")

def generate_realistic_timestamp(base_date=None, hours_offset=None, scenario_description='', is_urgent=False):
    """
    Generates timestamps that cluster around business hours with scenario-aware patterns.

    Args:
        base_date: Starting date/time
        hours_offset: Hours to add from base_date
        scenario_description: Scenario type for context-aware timing
        is_urgent: Whether this is an urgent/time-sensitive communication
    """
    if base_date is None:
        base_date = datetime.now() - timedelta(days=random.randint(10, 100))

    # Determine timing pattern based on scenario
    is_fraud_scenario = any(keyword in scenario_description.lower()
                            for keyword in ['fraud', 'hiding', 'coverup', 'destruction', 'shred', 'manipulat'])
    is_privilege_scenario = 'privilege' in scenario_description.lower() or 'confidential' in scenario_description.lower()

    if hours_offset:
        if is_urgent or is_fraud_scenario:
            # Urgent/fraud scenarios: Quick replies (30 min to 4 hours)
            if is_fraud_scenario and random.random() < 0.3:
                # Some fraud emails show deliberate delays (let things cool down)
                new_date = base_date + timedelta(hours=random.randint(24, 72), minutes=random.randint(0, 59))
            else:
                # Most urgent replies are fast
                new_date = base_date + timedelta(minutes=random.randint(30, 240))
        elif is_privilege_scenario:
            # Legal consultations: Often same-day or next business day
            new_date = base_date + timedelta(hours=random.randint(2, 24), minutes=random.randint(0, 59))
        else:
            # Normal business email patterns
            new_date = base_date + timedelta(hours=hours_offset, minutes=random.randint(1, 59))
    else:
        new_date = base_date

    # Apply business hours (80% of emails during 8 AM - 6 PM)
    if random.random() < 0.8:
        new_date = new_date.replace(hour=random.randint(8, 18), minute=random.randint(0, 59))
    else:
        # 20% outside business hours
        if random.random() < 0.5:
            # Early morning (6-8 AM)
            new_date = new_date.replace(hour=random.randint(6, 7), minute=random.randint(0, 59))
        else:
            # Evening (6 PM - 11 PM) - fraud/urgent scenarios more likely
            hour_max = 23 if (is_fraud_scenario or is_urgent) else 21
            new_date = new_date.replace(hour=random.randint(18, hour_max), minute=random.randint(0, 59))

    # Handle weekend/Friday evening patterns
    if random.random() < 0.9:  # 90% respect weekends
        # Skip weekends
        while new_date.weekday() >= 5:
            new_date += timedelta(days=1)

        # Friday evening -> Monday morning pattern
        if new_date.weekday() == 4 and new_date.hour >= 17:  # Friday after 5 PM
            if random.random() < 0.7:  # 70% chance it waits until Monday
                # Jump to Monday morning
                days_to_add = 7 - new_date.weekday()  # Days until Monday
                new_date = new_date + timedelta(days=days_to_add)
                new_date = new_date.replace(hour=random.randint(8, 10), minute=random.randint(0, 59))

    return new_date

# --- Core Generation Functions ---

def generate_email_thread(prompts, base_filename, output_dir, context_block, variables, personnel_map, scenario_description, attachment_config, near_dup_prob, run_count=1, is_noise=False, stats=None, config_temp=None):
    """Generates a threaded email conversation."""
    previous_message_id, references, generated_count = None, [], 0
    previous_email_content, previous_email_date = None, None
    temperature = get_temperature_for_scenario('thread', is_noise, config_temp)
    for i, prompt_obj in enumerate(prompts):
        probability = prompt_obj.get('probability', 1.0)
        if random.random() > probability:
            print(f"  ... Skipping a prompt in thread based on probability < {probability}")
            continue

        randomized_prompt = get_randomized_prompt(prompt_obj, variables, personnel_map, run_count)
        sender_name = get_sender_name_from_prompt(randomized_prompt, personnel_map)
        style_instruction = ""
        if sender_name and personnel_map.get(sender_name, {}).get('style'):
            style = personnel_map[sender_name]['style']
            style_instruction = f"\n\nIMPORTANT: Write this email in the style of {sender_name}: {style}"
        if previous_email_content:
            quoted_body = format_quoted_body(previous_email_content, previous_email_date)
            full_prompt = f"{context_block}\n\nYou are drafting a reply to the following email:\n\n---\n{quoted_body}\n---\n\nYour task: {randomized_prompt}{style_instruction}"
        else:
            full_prompt = f"{context_block}\n\nTask: {randomized_prompt}{style_instruction}"
        email_content = generate_email_content_from_llm(full_prompt, temperature)
        if not email_content: continue

        if random.random() < near_dup_prob:
            print("  -> Creating a near-duplicate variation...")
            email_content = create_near_duplicate(email_content)

        if previous_email_content and not email_content.get('subject', '').lower().startswith('re:'):
            email_content['subject'] = f"Re: {previous_email_content.get('subject', '')}"
        if previous_email_content:
            email_content['body'] += format_quoted_body(previous_email_content, previous_email_date)
        
        sender_email = email_content.get('sender_email', '')
        if '@' in sender_email:
            domain = sender_email.split('@')[1]
        else:
            domain = 'synthetic.local' # Safe fallback domain
            print(f"  !!! WARNING: LLM returned invalid sender_email: '{sender_email}'. Using fallback domain.")
        
        current_message_id = f"<{uuid.uuid4()}@{domain}>"
        headers = {'Message-ID': current_message_id}

        if previous_message_id:
            headers['In-Reply-To'] = previous_message_id
            headers['References'] = " ".join(references)
        references.append(current_message_id)

        # Detect urgency from email content for better timestamp realism
        is_urgent = any(keyword in email_content.get('subject', '').lower() + email_content.get('body', '').lower()
                       for keyword in ['urgent', 'asap', 'immediately', 'critical', 'emergency', 'catastrophic'])

        current_email_date = generate_realistic_timestamp(
            previous_email_date,
            random.randint(1, 48) if previous_email_date else None,
            scenario_description=scenario_description,
            is_urgent=is_urgent
        )

        dynamic_base_filename = f"{base_filename}_{generated_count + 1}"
        create_and_save_email(dynamic_base_filename, email_content, output_dir, current_email_date, personnel_map, scenario_description, attachment_config, headers, stats)

        generated_count += 1
        previous_message_id, previous_email_content, previous_email_date = current_message_id, email_content, current_email_date
    return generated_count

def generate_standalone_email(prompt_template, base_filename, output_dir, context_block, variables, personnel_map, scenario_description, attachment_config, near_dup_prob, run_count=1, is_noise=False, stats=None, config_temp=None):
    """Generates a standalone email."""
    randomized_prompt = get_randomized_prompt(prompt_template, variables, personnel_map, run_count)
    sender_name = get_sender_name_from_prompt(randomized_prompt, personnel_map)
    style_instruction = ""
    if sender_name and personnel_map.get(sender_name, {}).get('style'):
        style = personnel_map[sender_name]['style']
        style_instruction = f"\n\nIMPORTANT: Write this email in the style of {sender_name}: {style}"
    full_prompt = f"{context_block}\n\nTask: {randomized_prompt}{style_instruction}"
    temperature = get_temperature_for_scenario('standalone', is_noise, config_temp)
    email_content = generate_email_content_from_llm(full_prompt, temperature)
    if not email_content: return 0

    if random.random() < near_dup_prob:
        print("  -> Creating a near-duplicate variation...")
        email_content = create_near_duplicate(email_content)

    # --- STRESS TEST: CHECK FOR BLAST EMAIL SCENARIO ---
    if "blast_email" in base_filename.lower():
        email_content = inject_blast_recipients(email_content, personnel_map)
    # ---------------------------------------------------

    sender_email = email_content.get('sender_email', '')
    if '@' in sender_email:
        domain = sender_email.split('@')[1]
    else:
        domain = 'synthetic.local' # Safe fallback domain
        print(f"  !!! WARNING: LLM returned invalid sender_email: '{sender_email}'. Using fallback domain.")
    
    headers = {'Message-ID': f"<{uuid.uuid4()}@{domain}>"}

    # Detect urgency for standalone emails too
    is_urgent = any(keyword in email_content.get('subject', '').lower() + email_content.get('body', '').lower()
                   for keyword in ['urgent', 'asap', 'immediately', 'critical', 'emergency', 'catastrophic'])

    email_date = generate_realistic_timestamp(
        scenario_description=scenario_description,
        is_urgent=is_urgent
    )

    create_and_save_email(base_filename, email_content, output_dir, email_date, personnel_map, scenario_description, attachment_config, headers, stats)
    return 1

def generate_calendar_event(prompt_template, base_filename, output_dir, context_block, variables, personnel_map, run_count=1, is_noise=False, stats=None, config_temp=None):
    """Generates a standalone .ics calendar event."""
    randomized_prompt = get_randomized_prompt(prompt_template, variables, personnel_map, run_count)
    full_prompt = f"{context_block}\n\nTask: {randomized_prompt}"
    temperature = get_temperature_for_scenario('calendar', is_noise, config_temp)
    event_content = generate_calendar_content_from_llm(full_prompt, temperature)
    if not event_content: return 0
    event_date = generate_realistic_timestamp()
    filename = f"{base_filename}.ics"
    create_and_save_calendar_event(filename, event_content, output_dir, event_date, stats)
    return 1

def generate_chat_scenario(prompts, base_filename, output_dir, context_block, variables, personnel_map, chat_format='slack', run_count=1, is_noise=False, stats=None, config_temp=None):
    """Orchestrates the creation of a chat/RSMF file."""
    prompt_obj = prompts[0]
    randomized_prompt = get_randomized_prompt(prompt_obj, variables, personnel_map, run_count)

    full_prompt = f"{context_block}\n\nTask: {randomized_prompt}\n\nGenerate a conversation history between these participants."

    chat_content = generate_chat_content_from_llm(full_prompt, get_temperature_for_scenario('chat', is_noise, config_temp))
    
    if not chat_content: return 0

    start_date = generate_realistic_timestamp()
    
    # 1. Slack
    if chat_format in ['slack', 'all']:
        create_and_save_slack_native(base_filename, chat_content, output_dir, start_date, personnel_map, stats)

    # 2. Teams (RSMF)
    if chat_format in ['teams', 'all']:
        create_and_save_rsmf(base_filename, chat_content, output_dir, start_date, personnel_map, stats)

    # 3. Webex (API Format)
    if chat_format in ['webex', 'all']:
        create_and_save_webex_native(base_filename, chat_content, output_dir, start_date, personnel_map, stats)

    return 1

def create_nested_containers(output_dir):
    """
    Stress Test Post-Processing:
    Zips individual custodian folders, then puts them all in a .tar.gz archive.
    """
    print("\n--- Starting Post-Processing: Nested Container Creation ---")
    
    zip_files = []
    
    # 1. Walk immediate subdirectories (Custodians)
    for item in os.listdir(output_dir):
        full_path = os.path.join(output_dir, item)
        if os.path.isdir(full_path):
            print(f"  > Zipping folder: {item}")
            try:
                shutil.make_archive(full_path, 'zip', full_path)
                zip_files.append(f"{full_path}.zip")
            except Exception as e:
                print(f"!!! Error zipping {item}: {e}")

    # 2. Create the Master TarGz
    tar_name = os.path.join(output_dir, f"Dataset_Nested_Export_{uuid.uuid4().hex[:6]}.tar.gz")
    print(f"  > Creating Master Archive: {tar_name}")
    
    try:
        with tarfile.open(tar_name, "w:gz") as tar:
            for zip_file in zip_files:
                tar.add(zip_file, arcname=os.path.basename(zip_file))
                
        # 3. Cleanup: Remove the intermediate zip files
        print("  > Cleaning up intermediate .zip files...")
        for zip_file in zip_files:
            if os.path.exists(zip_file):
                os.remove(zip_file)
                
        print(f"--- Nested Container Created Successfully: {tar_name} ---")

    except Exception as e:
        print(f"!!! Error creating master archive: {e}")

def generate_protocol_document(output_dir, scenario_filter, config):
    """
    Generates an e-discovery investigation protocol document based on the scenario type.
    """
    print("\n--- Generating E-Discovery Protocol Document ---")

    # Determine scenario type from filter
    investigation_type = scenario_filter if scenario_filter and scenario_filter != 'all' else 'mixed'

    # Protocol templates by investigation type
    protocols = {
        'antitrust': {
            'title': 'Antitrust Price-Fixing Investigation',
            'matter_name': 'ACME Inc. v. Phoney Tunes Competition Investigation',
            'description': 'Investigation into potential horizontal price-fixing between ACME Inc. and competitor Phoney Tunes Inc.',
            'key_issues': [
                'Horizontal price-fixing agreements',
                'Market allocation schemes',
                'Competitor coordination on pricing',
                'Executive knowledge or approval'
            ],
            'custodians': [
                'Jamie Chen (ACME Regional Sales Director) - Key coordinator',
                'Rachel Quinn (Phoney Tunes Senior Account Executive) - Competitor contact',
                'Anthony Brooks (Phoney Tunes Account Executive) - Competitor contact',
                'Taylor Brooks (ACME CEO) - Potential executive knowledge',
                'Casey Bennett (ACME CFO) - BCCed on updates'
            ],
            'search_terms': [
                '"shared game plan"', '"ticket prices"', '"pricing strategy"',
                '"15% baseline"', '"our arrangement"', '"holding the line"',
                '"stabilize" AND "market"', '"Phoney Tunes" AND ("pricing" OR "agreement")',
                'competitor* w/10 (pricing OR price OR coordination)',
                'collu* OR conspir*'
            ],
            'date_range': 'July 1, 2024 - December 31, 2024',
            'expected_hot_docs': '7-10 emails showing collusion, 1 calendar invite, 1-2 chat logs',
            'privilege_concerns': 'GC communications if internal investigation started'
        },
        'safety_fraud': {
            'title': 'Product Safety Fraud Investigation',
            'matter_name': 'Project Phoenix Safety Test Manipulation Investigation',
            'description': 'Investigation into potential falsification of safety test results and evidence destruction for Project Phoenix.',
            'key_issues': [
                'Safety test result manipulation',
                'Evidence destruction',
                'False statements to clients',
                'Regulatory compliance violations'
            ],
            'custodians': [
                'Casey Mitchell (R&D Lead Engineer) - Orchestrator',
                'Peyton Parker (QA Specialist) - Initial discoverer, then complicit',
                'Taylor Brooks (CEO) - CCed on destruction order',
                'Rebecca Turner (Coyote Industries) - Client receiving false data',
                'Drew Foster (General Counsel) - Potential privilege issues'
            ],
            'search_terms': [
                '"catastrophic failure"', '"test results"', '"shred" OR "destroy"',
                '"re-calibrate" OR "recalibrate"', '"sensor data"', '"outlier*"',
                '"Project Phoenix" AND (test OR failure OR report)',
                'falsif* OR manipulat* OR conceal*',
                '"cannot leave this room"', '"no traces"'
            ],
            'date_range': 'September 1, 2024 - November 30, 2024',
            'expected_hot_docs': '6-8 emails showing coverup, test reports (attachments)',
            'privilege_concerns': 'Attorney-client communications about disclosure obligations'
        },
        'hr_misconduct': {
            'title': 'Workplace Harassment Investigation',
            'matter_name': 'ACME Inc. HR Investigation - Case #2025-003',
            'description': 'Investigation into workplace harassment allegations against Jamie Chen.',
            'key_issues': [
                'Workplace harassment complaints',
                'Inappropriate conduct toward subordinate',
                'Hostile work environment',
                'Witness corroboration'
            ],
            'custodians': [
                'Peyton Parker (QA Specialist) - Complainant',
                'Drew Foster (General Counsel) - Investigation lead',
                'Jamie Chen (Regional Sales Director) - Accused',
                'Casey Mitchell (R&D Lead Engineer) - Witness',
                'Taylor Brooks (CEO) - Investigation approval'
            ],
            'search_terms': [
                '"inappropriate"', '"uncomfortable"', '"harassment"',
                '"HR matter"', '"investigation"', '"confidential"',
                '"witness" AND ("statement" OR "corroborate")',
                'complain* OR allegat*',
                '"policy violation"'
            ],
            'date_range': 'July 1, 2024 - November 30, 2024',
            'expected_hot_docs': '8-10 emails/chats showing misconduct and investigation',
            'privilege_concerns': 'GC communications with complainant and CEO'
        }
    }

    # Get the appropriate protocol or use generic for 'all'
    if investigation_type in protocols:
        protocol = protocols[investigation_type]
    else:
        # Generic protocol for 'all' scenario
        protocol = {
            'title': 'Multi-Issue Investigation',
            'matter_name': 'ACME Inc. Comprehensive Investigation',
            'description': 'Comprehensive investigation covering multiple potential issues.',
            'key_issues': ['Multiple investigation types - see individual scenarios'],
            'custodians': ['All ACME personnel', 'External contacts as relevant'],
            'search_terms': ['Varies by scenario - see focused protocols'],
            'date_range': 'January 1, 2024 - December 31, 2024',
            'expected_hot_docs': 'Varies by scenario',
            'privilege_concerns': 'Attorney-client communications throughout'
        }

    # Build protocol document
    protocol_content = f"""# E-DISCOVERY INVESTIGATION PROTOCOL

## Matter Information

**Matter Name:** {protocol['matter_name']}
**Investigation Type:** {protocol['title']}
**Date Range:** {protocol['date_range']}
**Protocol Date:** {datetime.now().strftime('%B %d, %Y')}

---

## Investigation Overview

{protocol['description']}

### Key Issues to Investigate

"""

    for issue in protocol['key_issues']:
        protocol_content += f"- {issue}\n"

    protocol_content += f"""
---

## Custodian List

**Key Custodians:**

"""

    for custodian in protocol['custodians']:
        protocol_content += f"- {custodian}\n"

    protocol_content += f"""
---

## Search Terms Protocol

**Search Terms (Boolean):**

"""

    for term in protocol['search_terms']:
        protocol_content += f"- {term}\n"

    protocol_content += f"""
**Search Methodology:**
- Apply terms across email subject, body, and attachments
- Include threading analysis for related communications
- Review hits for responsiveness and privilege

---

## Expected Document Yield

**Hot Documents:** {protocol['expected_hot_docs']}

**Document Categories:**
1. **Responsive Hot Documents** - Direct evidence of key issues
2. **Supporting Documents** - Contextual or corroborating evidence
3. **Privileged Documents** - Attorney-client communications to be withheld
4. **Non-Responsive** - Business communications outside investigation scope

---

## Privilege Protocol

**Privilege Concerns:** {protocol['privilege_concerns']}

**Privilege Review Process:**
1. Identify communications with General Counsel (Drew Foster)
2. Identify communications with external counsel (Eleanor Vance, Wright & Lee LLP)
3. Apply attorney-client privilege analysis
4. Create privilege log for withheld documents
5. Second-level review for privilege assertions

**Privilege Designations:**
- **Privileged** - Withheld under attorney-client privilege
- **Privileged - Work Product** - Withheld under work product doctrine
- **Non-Privileged** - Produced

---

## Quality Control Measures

**Sampling Methodology:**
- Review 5% random sample of responsive documents
- 100% review of hot documents
- Privilege review of all attorney communications

**Review Standards:**
- Relevance to key issues
- Temporal relevance (within date range)
- Custodian relevance
- Privilege determination

---

## EAIDA Testing Metrics

**Success Criteria:**

1. **Precision:** Did EAIDA find the expected hot documents?
2. **Recall:** Did EAIDA miss any critical evidence?
3. **Privilege Accuracy:** Correctly identified privileged communications?
4. **Relationship Mapping:** Identified key players and connections?
5. **Timeline Accuracy:** Reconstructed sequence of events?

**Expected EAIDA Performance:**
- Hot document identification: 90%+
- Privilege identification: 85%+
- False positives: <10%
- Key custodian identification: 100%

---

## Generated Dataset Information

**Output Directory:** {output_dir}
**Scenario Filter:** {scenario_filter}
**Generation Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

**Dataset Contents:**
- Synthetic emails organized by custodian
- Chat logs (Slack/Teams/Webex format)
- Calendar invites
- Document attachments
- Privilege documents (automatically included)
- Contextual noise documents

---

## Notes for EAIDA Testing

1. **Load the entire dataset** into EAIDA for testing
2. **Use search terms** from this protocol to test search functionality
3. **Validate hot documents** against expected yields
4. **Test privilege screening** on GC and external counsel communications
5. **Assess relationship mapping** between key custodians
6. **Evaluate timeline reconstruction** capabilities

**This is synthetic data for testing purposes only.**

---

*End of Protocol Document*
"""

    # Save protocol document
    protocol_filename = os.path.join(output_dir, "INVESTIGATION_PROTOCOL.md")
    try:
        with open(protocol_filename, 'w', encoding='utf-8') as f:
            f.write(protocol_content)
        print(f"  > Protocol document created: {protocol_filename}")
        print("--- Protocol Document Generation Complete ---")
    except Exception as e:
        print(f"!!! Error creating protocol document: {e}")

# --- Main Orchestration Logic ---
if __name__ == "__main__":
    selected_config_file = select_config_file()
    if not selected_config_file: exit()
    target_item_count = get_target_email_count()
    
    # --- Prompt for Chat Format ---
    chat_format_pref = get_chat_format_preference()
    
    # --- Prompt for Log Size (Stress Test) ---
    log_size_mb = get_log_size_preference()
    
    # --- Prompt for Nested Container (Stress Test) ---
    create_container = get_container_preference()

    # --- Prompt for Protocol Document Generation ---
    generate_protocol = get_protocol_preference()

    config = load_config(selected_config_file)
    if not config: exit()
    print("\nConfiguration loaded.")

    # Check if user selected config-acme.yaml and prompt for scenario filter
    if selected_config_file == 'config-acme.yaml':
        print("\n" + "="*80)
        print("You selected config-acme.yaml (Master Configuration with All Scenarios)")
        print("="*80)
        print("\nThis config contains ALL investigation types mixed together.")
        print("For realistic investigation testing, you should focus on ONE investigation type.")

        user_filter_choice = get_scenario_filter_preference()

        # Override the config's scenario_filter with user's choice
        config['general_settings']['scenario_filter'] = user_filter_choice

        print("\nApplying filter to dataset generation...")

    print("\nStarting large-scale item generation...")

    output_dir = config['general_settings']['output_directory']
    if os.path.exists(output_dir):
        print(f"Warning: Output directory '{output_dir}' already exists.")

    # Apply scenario filtering
    scenario_filter = config.get('general_settings', {}).get('scenario_filter', 'all')
    all_scenarios = config['scenarios']
    filtered_scenarios = filter_scenarios_by_type(all_scenarios, scenario_filter)

    if scenario_filter and scenario_filter != 'all':
        print(f"\n[SCENARIO FILTER ACTIVE]: '{scenario_filter}'")
        print(f"  Total scenarios in config: {len(all_scenarios)}")
        print(f"  Scenarios after filter: {len(filtered_scenarios)}")
        if len(filtered_scenarios) == 0:
            print("\n!!! WARNING: No scenarios match the filter! Check your 'scenario_filter' setting.")
            print("    Valid options: 'all', 'antitrust', 'safety_fraud', 'hr_misconduct'")
            print("    Note: 'legal_privilege' removed - privilege scenarios (S3) now included with all investigations")
            exit()

    # Replace the scenarios list with the filtered version
    config['scenarios'] = filtered_scenarios

    context = build_context_block(config['company_profiles'])
    personnel_map = build_personnel_map(config['company_profiles'])
    attachment_config = config.get('attachments', {})
    
    # Pass the log size into the config dictionary so it reaches the email function
    if attachment_config:
        attachment_config['log_size_mb'] = log_size_mb
    
    # --- STATS TRACKING INIT ---
    stats = {
        'emails': 0,
        'calendar_events': 0,
        'rsmf_chats': 0,
        'attachments': 0,
        'attachment_types': {},
        'scenarios_triggered': {},
        'stress_tests_triggered': [],
        'email_dates': [],  # Track all email dates for date range calculation
        'custodians': set()  # Track unique custodian emails
    }

    scenario_run_counts = {}
    generated_item_count, run_counter = 0, 1

    # Thread-safe locks for stats and counters
    stats_lock = threading.Lock()
    count_lock = threading.Lock()

    # Configuration: Number of parallel workers (tune based on your Azure OpenAI TPM limits)
    MAX_WORKERS = 10

    def process_scenario_worker(scenario, run_counter, scenario_run_counts_local):
        """
        Worker function to process a single scenario in parallel.
        Returns tuple: (items_created, scenario_desc, scenario_updates)
        """
        try:
            scenario_id = scenario['base_filename']
            current_run = scenario_run_counts_local.get(scenario_id, 0) + 1

            dynamic_base_filename = f"{scenario['base_filename']}_r{run_counter}_{uuid.uuid4().hex[:6]}"
            scenario_desc = scenario['description']

            print(f"  Running Scenario: {scenario_desc} (Occurrence #{current_run})")

            variables = scenario.get('prompt_variables', None)
            near_dup_prob = scenario.get('near_duplicate_probability', 0.0)
            items_created = 0
            is_noise = 'noise' in scenario['base_filename'].lower()
            prompts_list = scenario['prompts']

            # Extract temperature from llm_settings if present
            config_temp = None
            if 'llm_settings' in scenario and 'temperature' in scenario['llm_settings']:
                config_temp = scenario['llm_settings']['temperature']

            # Track stress tests
            stress_test_triggered = None
            if "blast_email" in scenario['base_filename']:
                stress_test_triggered = "Blast Email Expansion"

            # Generate based on scenario type
            if scenario['type'] == 'thread':
                items_created = generate_email_thread(prompts_list, dynamic_base_filename, output_dir, context, variables, personnel_map, scenario_desc, attachment_config, near_dup_prob, current_run, is_noise, stats, config_temp)
            elif scenario['type'] == 'standalone':
                items_created = generate_standalone_email(prompts_list[0], dynamic_base_filename, output_dir, context, variables, personnel_map, scenario_desc, attachment_config, near_dup_prob, current_run, is_noise, stats, config_temp)
            elif scenario['type'] == 'calendar_event':
                items_created = generate_calendar_event(prompts_list[0], dynamic_base_filename, output_dir, context, variables, personnel_map, current_run, is_noise, stats, config_temp)
            elif scenario['type'] == 'chat':
                items_created = generate_chat_scenario(prompts_list, dynamic_base_filename, output_dir, context, variables, personnel_map, chat_format_pref, current_run, is_noise, stats, config_temp)

            if items_created > 0:
                print(f"  > Generated {items_created} item(s) for this scenario.")

            return (items_created, scenario_desc, scenario_id, stress_test_triggered)

        except Exception as e:
            print(f"  !!! ERROR processing scenario {scenario.get('description', 'unknown')}: {e}")
            return (0, scenario.get('description', 'unknown'), scenario.get('base_filename', 'unknown'), None)

    while generated_item_count < target_item_count:
        print(f"\n--- Starting Generation Run #{run_counter} ---")
        scenarios = config['scenarios'].copy()
        random.shuffle(scenarios)

        # Build list of scenarios to process in this run
        scenarios_to_process = []
        for scenario in scenarios:
            if generated_item_count >= target_item_count:
                break
            scenarios_to_process.append(scenario)

        # Process scenarios in parallel with ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            # Create local copy of run counts for this batch
            scenario_run_counts_local = scenario_run_counts.copy()

            # Submit all scenarios to thread pool
            future_to_scenario = {}
            for scenario in scenarios_to_process:
                # Update local run count
                scenario_id = scenario['base_filename']
                scenario_run_counts_local[scenario_id] = scenario_run_counts_local.get(scenario_id, 0) + 1

                # Submit to thread pool
                future = executor.submit(process_scenario_worker, scenario, run_counter, scenario_run_counts_local)
                future_to_scenario[future] = scenario

            # Collect results as they complete
            for future in as_completed(future_to_scenario):
                items_created, scenario_desc, scenario_id, stress_test = future.result()

                # Thread-safe updates to shared state
                with stats_lock:
                    stats['scenarios_triggered'][scenario_desc] = stats['scenarios_triggered'].get(scenario_desc, 0) + 1
                    if stress_test and stress_test not in stats['stress_tests_triggered']:
                        stats['stress_tests_triggered'].append(stress_test)

                with count_lock:
                    scenario_run_counts[scenario_id] = scenario_run_counts.get(scenario_id, 0) + 1
                    if items_created > 0:
                        generated_item_count += items_created
                        print(f"  > Progress: {generated_item_count} / {target_item_count} total items generated.")

        run_counter += 1

    # --- POST PROCESSING: NESTED CONTAINER ---
    if create_container:
        create_nested_containers(output_dir)
        stats['stress_tests_triggered'].append("Recursive Containerization")

    # --- REPORT GENERATION LOGIC ---
    total_docs = stats['emails'] + stats['calendar_events'] + stats.get('rsmf_chats', 0) + stats['attachments']
    
    # Categorize Scenarios for the Report
    # Signal scenarios include investigation-specific tags: S1, S2, S_HR, and always S3 (privilege)
    signal_scenarios = {k:v for k,v in stats['scenarios_triggered'].items() if "(S1)" in k or "(S2)" in k or "(S3)" in k or "(S_HR)" in k}
    noise_scenarios = {k:v for k,v in stats['scenarios_triggered'].items() if k not in signal_scenarios}

    # Build dynamic signal description based on what's actually in the dataset
    signal_types = []
    if any("(S1)" in k for k in signal_scenarios.keys()):
        signal_types.append("Price-Fixing (Antitrust)")
    if any("(S2)" in k for k in signal_scenarios.keys()):
        signal_types.append("Safety Fraud (Product Liability)")
    if any("(S_HR)" in k for k in signal_scenarios.keys()):
        signal_types.append("Workplace Harassment (HR)")
    if any("(S3)" in k for k in signal_scenarios.keys()):
        signal_types.append("Privilege (Attorney-Client)")

    signal_description = ", ".join(signal_types) if signal_types else "Various investigation scenarios"

    # Build dynamic noise description based on what's actually in the dataset
    noise_types = []
    if any("sales" in k.lower() for k in noise_scenarios.keys()):
        noise_types.append("Sales chatter")
    if any("hr" in k.lower() or "admin" in k.lower() for k in noise_scenarios.keys()):
        noise_types.append("HR announcements")
    if any("project" in k.lower() or "engineering" in k.lower() for k in noise_scenarios.keys()):
        noise_types.append("IT/Project emails")
    if any("personal" in k.lower() for k in noise_scenarios.keys()):
        noise_types.append("Personal emails")
    if any("meeting" in k.lower() or "scheduling" in k.lower() for k in noise_scenarios.keys()):
        noise_types.append("Meeting scheduling")
    if any("typo" in k.lower() or "fragment" in k.lower() or "vague" in k.lower() for k in noise_scenarios.keys()):
        noise_types.append("Low-quality emails")

    noise_description = ", ".join(noise_types) if noise_types else "Business noise"

    print("\n" + "="*80)
    print(f"   SYNTHETIC DATASET CERTIFICATION REPORT   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("="*80)

    # Calculate date range and custodian statistics
    date_range_str = "N/A"
    days_span = 0
    if stats['email_dates']:
        min_date = min(stats['email_dates'])
        max_date = max(stats['email_dates'])
        days_span = (max_date - min_date).days + 1
        date_range_str = f"{min_date.strftime('%b %d, %Y')} - {max_date.strftime('%b %d, %Y')} ({days_span} days)"

    # Count core vs expansion custodians (assuming core 5 have specific patterns in their names/roles)
    # For now, just show total active custodians
    total_custodians = len(stats['custodians'])

    print(f"\n[1] DATASET VOLUME & DIVERSITY")
    print(f"    Total Readable Documents: {total_docs}")
    print(f"    Simulated Date Range:     {date_range_str}")
    print(f"    Active Custodians:        {total_custodians}")
    print(f"    ---------------------------------------")
    print(f"    • Emails (.eml):             {stats['emails']:<5} (RFC-compliant, headers included)")
    print(f"    • Chats (Slack/Teams/Webex): {stats.get('rsmf_chats', 0):<5} (Modern short-message format)")
    print(f"    • Calendar (.ics):           {stats['calendar_events']:<5} (Meeting invites)")
    print(f"    • Attachments:               {stats['attachments']:<5} (Context-aware content)")

    if stats['attachment_types']:
        print(f"      - File Types Generated: {', '.join([k for k in stats['attachment_types'].keys()])}")

    print(f"\n[2] TRAINING & VALIDATION VALUE (Signal vs. Noise)")
    print(f"    This dataset creates a 'Needle in a Haystack' environment for reviewer training.")
    print(f"    ---------------------------------------")
    print(f"    ► THE NEEDLES (Critical Evidence): {sum(signal_scenarios.values())} events")
    print(f"      Includes: {signal_description}")
    print(f"    ► THE HAYSTACK (Contextual Noise): {sum(noise_scenarios.values())} events")
    print(f"      Includes: {noise_description}")
    
    print(f"\n[3] TECHNICAL STRESS TEST ARTIFACTS")
    print(f"    This run included specific challenges to test ingestion and processing stability.")
    print(f"    ---------------------------------------")
    
    # Dynamic reporting on stress tests
    has_stress = False
    
    # 1. Check for Large Logs
    log_count = stats['attachment_types'].get('.log', 0)
    if log_count > 0:
        has_stress = True
        print(f"    [X] Throughput Test:      Generated {log_count} System Log files.")
        if log_size_mb >= 50:
            print(f"                              - Included Jumbo Files ({log_size_mb} MB) to test timeout/memory.")
    
    # 2. Check for Blast Email
    if "Blast Email Expansion" in stats['stress_tests_triggered']:
        has_stress = True
        print(f"    [X] Metadata Explosion:   Generated 'Blast Emails' with 500+ recipients in To/CC header.")
    
    # 3. Check for Container
    if create_container:
        has_stress = True
        print(f"    [X] Recursion Depth:      Wrapped output in Nested Container (TarGz -> Zip -> Files).")
    
    # 4. Check for Foreign Language
    if any("Foreign language" in k for k in noise_scenarios):
        has_stress = True
        print(f"    [X] Language Detection:   Included non-English business emails (CJK/Euro).")

    if not has_stress:
        print(f"    [ ] No specific stress-test artifacts were triggered in this run.")

    # Generate protocol document if requested
    if generate_protocol:
        generate_protocol_document(output_dir, scenario_filter, config)

    print(f"\n[4] OUTPUT LOCATION")
    print(f"    Directory: {os.path.abspath(output_dir)}")
    if create_container:
        print(f"    Archive:   Dataset_Nested_Export_*.tar.gz")
    if generate_protocol:
        print(f"    Protocol:  INVESTIGATION_PROTOCOL.md")
    print("="*80)